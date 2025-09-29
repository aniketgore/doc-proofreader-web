# Copyright caerulex 2024

"""Main entrypoint of doc-proofreader."""

from datetime import datetime
from docx import Document
from pathlib import Path
from dotenv import load_dotenv
import os
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading
from doc_proofreader.prompts.system_prompts import DEFAULT_SYSTEM_PROMPT
from doc_proofreader.prompts.user_prompts import USER_PROMPT
from doc_proofreader.llm.client_factory import ClientFactory
from doc_proofreader.chunk_utils import parse_chunk_size, get_optimal_chunk_size, validate_chunk_size

load_dotenv()


def docx_to_formatted_text(file_path):
    doc = Document(file_path)
    formatted_text = ""

    for para in doc.paragraphs:
        for run in para.runs:
            text = run.text
            if run.bold and run.italic:
                formatted_text += "<b><i>" + text + "</i></b>"
            elif run.bold:
                formatted_text += "<b>" + text + "</b>"
            elif run.italic:
                formatted_text += "<i>" + text + "</i>"
            else:
                formatted_text += text
        formatted_text += "\n"  # Add a newline after each paragraph for readability

    return formatted_text


def chunk_text(text, chunk_size=2000):  # 12000
    # This function is aware of the HTML-like tags to avoid splitting them.
    # TODO(caerulex): Don't split up a sentence between two chunks.
    chunks = []
    current_chunk = ""
    words = text.split(" ")
    for word in words:
        if len(current_chunk) + len(word) < chunk_size:
            current_chunk += word + " "
        else:
            chunks.append(current_chunk)
            current_chunk = word + " "
    if current_chunk:
        chunks.append(current_chunk)
    return chunks


def docx_to_chunks(file_path, chunk_size):
    doc = Document(file_path)

    chunks = []
    current_paragraph = ""
    current_chunk = ""

    for para in doc.paragraphs:
        current_paragraph = ""
        for run in para.runs:
            text = run.text
            if run.bold and run.italic:
                current_paragraph += "<b><i>" + text + "</i></b>"
            elif run.bold:
                current_paragraph += "<b>" + text + "</b>"
            elif run.italic:
                current_paragraph += "<i>" + text + "</i>"
            else:
                current_paragraph += text
        current_paragraph += (
            "  \n"  # Add a newline after each paragraph for readability
        )
        current_chunk += current_paragraph

        # Check if the current paragraph is too long and chunk it if necessary
        if len(current_chunk) > chunk_size:  # Adjust the chunk size as needed
            # If the paragraph is not too long, append it to the list
            chunks.append(current_chunk)
            current_chunk = ""

    # Add the remainder that is smaller than the chunk size.
    if current_chunk:
        chunks.append(current_chunk)

    return chunks


def process_chunk_with_llm(
    chunk: str, additional_instructions: str, client, model: str = None, chunk_index: int = 0
):
    thread_id = threading.current_thread().ident
    print(f"Processing chunk {chunk_index + 1} (thread {thread_id})...")
    messages = [
        {"role": "system", "content": DEFAULT_SYSTEM_PROMPT},
        {"role": "user", "content": USER_PROMPT},
        {"role": "user", "content": chunk},
    ]
    if additional_instructions:
        messages.insert(2, {"role": "user", "content": additional_instructions})

    result = client.create_completion(
        messages=messages,
        model=model,
        temperature=0.2,
    )
    if result.strip() == "No errors were found.":
        result = ""
    print(f"✅ Chunk {chunk_index + 1} completed")
    return result


def aggregate_outputs(outputs: list[str]):
    # Simple aggregation - join all outputs.
    return " ".join(outputs)


def export_proofreads(document_path: str, output: str, save_dir: str):
    date = datetime.now().strftime("%m.%d.%Y_%H.%M.%S")
    name = Path(document_path).stem

    # Create the full path for the output file
    output_file_name = f"results_{name}_{date}.txt"
    if save_dir:
        save_path = Path(save_dir)
        if not save_path.exists() or not save_path.is_dir():
            # Create the directory if it does not exist
            save_path.mkdir(parents=True, exist_ok=True)
        output_file_path = save_path / output_file_name
    else:
        output_file_path = output_file_name

    with open(output_file_path, "w", encoding="utf-8") as file:
        file.write(output)
    print(f"Results saved to '{output_file_path}'")


def proofread_document(
    document_path: str,
    additional_instructions: str = "",
    save_outputs: bool = True,
    save_dir: str = "",
    provider: str = None,
    model: str = None,
    estimate_cost: bool = False,
    chunk_size_arg: str = None,
    parallel: bool = True,
    max_workers: int = 3,
) -> str:
    # Create LLM client
    client = ClientFactory.create_client(provider=provider, model_name=model)
    model_info = client.get_model_info()

    # Determine chunk size
    if chunk_size_arg:
        if chunk_size_arg.lower() == 'auto':
            chunk_size = get_optimal_chunk_size(model_info['name'], model_info['context_window'])
            print(f"🤖 Auto chunk size: {chunk_size:,} chars (~{chunk_size//5.5:.0f} words) for {model_info['name']}")
        else:
            chunk_size = parse_chunk_size(chunk_size_arg)
            is_valid, warning = validate_chunk_size(chunk_size, model_info['name'], model_info['context_window'])
            if not is_valid:
                print(f"❌ {warning}")
                return ""
            if warning:
                print(warning)
    else:
        chunk_size = 27500  # Default 5000 words

    # Get document content and estimate cost if requested
    if estimate_cost:
        doc_content = docx_to_formatted_text(document_path)
        cost = client.estimate_cost(doc_content)
        print(f"\n📊 Cost Estimation:")
        print(f"  Model: {model_info['name']} ({model_info['provider']})")
        print(f"  Estimated cost: ${cost:.4f}")
        print(f"  Context window: {model_info['context_window']:,} tokens")
        print(f"  Chunk size: {chunk_size:,} chars (~{chunk_size//5.5:.0f} words)")
        response = input("\nProceed with proofreading? (y/n): ")
        if response.lower() != 'y':
            print("Proofreading cancelled.")
            return ""

    chunks = docx_to_chunks(document_path, chunk_size)
    print(f"📄 Document split into {len(chunks)} chunks")
    print(f"⚡ Processing {'in parallel' if parallel and len(chunks) > 1 else 'sequentially'}...")

    if parallel and len(chunks) > 1:
        # Parallel processing
        results = [None] * len(chunks)  # Preserve order
        with ThreadPoolExecutor(max_workers=min(max_workers, len(chunks))) as executor:
            # Submit all chunks
            future_to_index = {
                executor.submit(
                    process_chunk_with_llm,
                    chunks[i],
                    additional_instructions,
                    ClientFactory.create_client(provider=provider, model_name=model),  # Fresh client per thread
                    model,
                    i
                ): i for i in range(len(chunks))
            }

            # Collect results as they complete
            for future in as_completed(future_to_index):
                index = future_to_index[future]
                try:
                    results[index] = future.result()
                except Exception as exc:
                    print(f"❌ Chunk {index + 1} generated an exception: {exc}")
                    results[index] = f"Error processing chunk {index + 1}: {exc}"

        aggregated_output = aggregate_outputs(results)
    else:
        # Sequential processing (fallback or single chunk)
        results = [
            process_chunk_with_llm(chunks[i], additional_instructions, client, model, i)
            for i in range(len(chunks))
        ]
        aggregated_output = aggregate_outputs(results)

    # Save the results to a text file
    if save_outputs:
        export_proofreads(document_path, aggregated_output, save_dir)

    return aggregated_output
