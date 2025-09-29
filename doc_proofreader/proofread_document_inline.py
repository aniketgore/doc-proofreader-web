# Copyright caerulex 2025

"""Doc-proofreader with inline corrections.

Outputs document with track changes on Mac platforms. Windows/Linux not
currently supported. Instead, users can manually run a "Compare Documents" in
Microsoft Word to see track changes between original and corrected documents.
"""

from datetime import datetime
from docx import Document
from pathlib import Path
from dotenv import load_dotenv
import os
import re
import sys
import subprocess
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
from doc_proofreader.prompts.system_prompts import DIRECT_EDIT_SYSTEM_PROMPT
from doc_proofreader.llm.client_factory import ClientFactory
from doc_proofreader.chunk_utils import parse_chunk_size, get_optimal_chunk_size, validate_chunk_size

load_dotenv()
SUPPORTED_OS = ["darwin"]


def clear_all_paragraphs(document):
    """
    Clears all paragraphs from a python-docx Document object.
    """
    # Iterate through paragraphs in reverse order to avoid issues with index shifts
    for paragraph in reversed(document.paragraphs):
        p = paragraph._element
        p.getparent().remove(p)
        # Clear the internal references to the deleted paragraph
        paragraph._p = paragraph._element = None


def process_chunk_for_direct_edit(
    chunk: str, additional_instructions: str, client, model: str = None, chunk_index: int = 0
):
    """Process chunk and return corrected text directly."""
    thread_id = threading.current_thread().ident
    print(f"Processing chunk {chunk_index + 1} for inline edits (thread {thread_id})...")

    messages = [
        {"role": "system", "content": DIRECT_EDIT_SYSTEM_PROMPT},
        {"role": "user", "content": chunk},
    ]

    if additional_instructions:
        messages.insert(
            1,
            {
                "role": "user",
                "content": f"Additional instructions: {additional_instructions}",
            },
        )

    result = client.create_completion(
        messages=messages,
        model=model,
        temperature=0.1,
    )
    print(f"✅ Chunk {chunk_index + 1} inline processing completed")
    return result


def create_corrected_document_from_chunks(
    original_doc_path, original_chunks, corrected_chunks
):
    """Create a new document with corrections applied, preserving formatting."""
    original_doc = Document(original_doc_path)
    corrected_doc = Document(original_doc_path)
    clear_all_paragraphs(corrected_doc)

    # Process paragraphs
    para_idx = 0
    for chunk_idx, (original_chunk, corrected_chunk) in enumerate(
        zip(original_chunks, corrected_chunks)
    ):
        corrected_lines = corrected_chunk.split("  \n")

        for line_idx, corrected_line in enumerate(corrected_lines):
            if corrected_line.strip():  # Skip empty lines
                if para_idx < len(original_doc.paragraphs):
                    original_para = original_doc.paragraphs[para_idx]
                    new_para = corrected_doc.add_paragraph()

                    # Copy paragraph style
                    try:
                        new_para.style = original_para.style
                    except Exception:
                        pass

                    # Parse and apply formatting
                    apply_formatted_text_to_paragraph(new_para, corrected_line)
                    para_idx += 1
                else:
                    # Add any remaining paragraphs
                    new_para = corrected_doc.add_paragraph()
                    apply_formatted_text_to_paragraph(new_para, corrected_line)

    return corrected_doc


def apply_formatted_text_to_paragraph(paragraph, formatted_text):
    """Apply formatted text with HTML-like tags to a paragraph."""
    # Clear existing runs
    paragraph.clear()

    # Pattern to match formatting tags and text
    pattern = r"(<b><i>|<b>|<i>|</b></i>|</b>|</i>)"
    parts = re.split(pattern, formatted_text)

    current_bold = False
    current_italic = False

    for part in parts:
        if part == "<b>":
            current_bold = True
        elif part == "</b>":
            current_bold = False
        elif part == "<i>":
            current_italic = True
        elif part == "</i>":
            current_italic = False
        elif part == "<b><i>":
            current_bold = True
            current_italic = True
        elif part == "</b></i>" or part == "</i></b>":
            current_bold = False
            current_italic = False
        elif part and part not in ["<b><i>", "<b>", "<i>", "</b></i>", "</b>", "</i>"]:
            # This is actual text
            if part.strip():
                run = paragraph.add_run(part)
                run.bold = current_bold
                run.italic = current_italic


def compare_documents_with_applescript(original_path, corrected_path) -> None:
    """Execute AppleScript in a single comprehensive block to avoid variable scoping issues."""

    # Convert to absolute paths
    original_abs = str(Path(original_path).resolve())
    corrected_abs = str(Path(corrected_path).resolve())

    print(f"🔍 Original document: {original_abs}")
    print(f"🔍 Corrected document: {corrected_abs}")

    # Single comprehensive AppleScript that does everything in one block
    comprehensive_script = f"""set old to "{original_abs}"
set new to "{corrected_abs}"
tell application "Microsoft Word"
    activate
    open old
    compare active document path new as text
end tell
"""

    try:
        # Execute the comprehensive script
        subprocess.run(
            ["osascript", "-e", comprehensive_script],
            capture_output=True,
            text=True,
            timeout=300,
        )
    except subprocess.TimeoutExpired:
        print("❌ AppleScript execution timed out")
        return False
    except Exception as e:
        print(f"❌ Error executing AppleScript: {e}")
        return False


def proofread_document_with_track_changes_mac(
    document_path: str,
    additional_instructions: str = "",
    save_dir: str = "",
    provider: str = None,
    model: str = None,
    estimate_cost: bool = False,
    chunk_size_arg: str = None,
    parallel: bool = True,
    max_workers: int = 3,
) -> str:
    """Main function to proofread document and create track changes version on Mac."""

    print(f"Processing document: {document_path}")

    # Create LLM client
    client = ClientFactory.create_client(provider=provider, model_name=model)
    model_info = client.get_model_info()

    # Determine chunk size
    if chunk_size_arg:
        if chunk_size_arg.lower() == 'auto':
            chunk_size = get_optimal_chunk_size(model_info['name'], model_info['context_window'])
            print(f"🤖 Auto chunk size: {chunk_size:,} chars (~{chunk_size//5.5:.0f} words) for {model_info['name']}")
        else:
            chunk_size = parse_chunk_size(chunk_size_arg)
            is_valid, warning = validate_chunk_size(chunk_size, model_info['name'], model_info['context_window'])
            if not is_valid:
                print(f"❌ {warning}")
                return ""
            if warning:
                print(warning)
    else:
        chunk_size = 27500  # Default 5000 words

    # Get document content and estimate cost if requested
    if estimate_cost:
        # Read document to estimate cost
        doc = Document(document_path)
        doc_text = "\n".join([para.text for para in doc.paragraphs])
        cost = client.estimate_cost(doc_text)
        print(f"\n📊 Cost Estimation:")
        print(f"  Model: {model_info['name']} ({model_info['provider']})")
        print(f"  Estimated cost: ${cost:.4f}")
        print(f"  Context window: {model_info['context_window']:,} tokens")
        print(f"  Chunk size: {chunk_size:,} chars (~{chunk_size//5.5:.0f} words)")
        response = input("\nProceed with proofreading? (y/n): ")
        if response.lower() != 'y':
            print("Proofreading cancelled.")
            return ""

    # Use existing chunking function
    original_chunks = docx_to_chunks(document_path, chunk_size)
    print(f"Document split into {len(original_chunks)} chunks")

    print(f"📄 Document split into {len(original_chunks)} chunks")
    print(f"⚡ Processing {'in parallel' if parallel and len(original_chunks) > 1 else 'sequentially'}...")

    # Process each chunk
    if parallel and len(original_chunks) > 1:
        # Parallel processing
        corrected_chunks = [None] * len(original_chunks)  # Preserve order
        with ThreadPoolExecutor(max_workers=min(max_workers, len(original_chunks))) as executor:
            # Submit all chunks
            future_to_index = {
                executor.submit(
                    process_chunk_for_direct_edit,
                    original_chunks[i],
                    additional_instructions,
                    ClientFactory.create_client(provider=provider, model_name=model),  # Fresh client per thread
                    model,
                    i
                ): i for i in range(len(original_chunks))
            }

            # Collect results as they complete
            for future in as_completed(future_to_index):
                index = future_to_index[future]
                try:
                    corrected_chunks[index] = future.result()
                except Exception as exc:
                    print(f"❌ Chunk {index + 1} generated an exception: {exc}")
                    corrected_chunks[index] = original_chunks[index]  # Fallback to original
    else:
        # Sequential processing
        corrected_chunks = []
        for i, chunk in enumerate(original_chunks):
            print(f"Processing chunk {i+1}/{len(original_chunks)}")
            corrected_text = process_chunk_for_direct_edit(
                chunk, additional_instructions, client, model, i
            )
            corrected_chunks.append(corrected_text)

    # Create corrected document
    corrected_doc = create_corrected_document_from_chunks(
        document_path, original_chunks, corrected_chunks
    )

    # Set up save paths
    date = datetime.now().strftime("%m.%d.%Y_%H.%M.%S")
    name = Path(document_path).stem

    if save_dir:
        save_path = Path(save_dir)
        save_path.mkdir(parents=True, exist_ok=True)
    else:
        save_path = Path(document_path).parent

    # Save corrected version first
    corrected_path = save_path / f"{name}_corrected_{date}.docx"
    corrected_doc.save(str(corrected_path))
    print(f"Corrected document saved: {corrected_path}")

    print("🔄 Attempting to create track changes document...")
    if sys.platform not in SUPPORTED_OS:
        print(
            "This application does not support track changes output on "
            f"'{sys.platform}'. Supported operating systems are: "
            f"{', '.join(SUPPORTED_OS)}."
        )
    else:
        compare_documents_with_applescript(
            document_path,
            corrected_path,
        )

    return str(corrected_path)


# Integration with your existing code structure
def docx_to_chunks(file_path, chunk_size):
    """Chunk document into manageable pieces. Default chunk_size=27500 chars (~5000 words)."""
    doc = Document(file_path)

    chunks = []
    current_paragraph = ""
    current_chunk = ""

    for para in doc.paragraphs:
        current_paragraph = ""
        for run in para.runs:
            text = run.text
            if run.bold and run.italic:
                current_paragraph += "<b><i>" + text + "</i></b>"
            elif run.bold:
                current_paragraph += "<b>" + text + "</b>"
            elif run.italic:
                current_paragraph += "<i>" + text + "</i>"
            else:
                current_paragraph += text
        current_paragraph += (
            "  \n"  # Add a newline after each paragraph for readability
        )
        current_chunk += current_paragraph

        # Check if the current paragraph is too long and chunk it if necessary
        if len(current_chunk) > chunk_size:  # Default: 27500 chars (~5000 words)
            # If the paragraph is not too long, append it to the list
            chunks.append(current_chunk)
            current_chunk = ""

    # Add the remainder that is smaller than the chunk size.
    if current_chunk:
        chunks.append(current_chunk)

    return chunks
